from __future__ import annotations

from pathlib import Path

from app.core.exceptions import ValidationError
from app.core.logging import get_logger

log = get_logger(__name__)

# A resume page typically yields >200 chars of text. Below this threshold we
# assume the document is image-based and fall back to OCR.
_MIN_TEXT_THRESHOLD = 200


def extract_text(path: Path) -> str:
    """
    Extract text from a PDF or DOCX. Returns the raw text (whitespace preserved
    enough to keep line structure). Does NOT run OCR — that's a separate step.
    """
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix in (".docx", ".doc"):
        return _extract_docx(path)
    raise ValidationError(f"Unsupported resume format: {suffix}. Use PDF or DOCX.")


def looks_image_based(text: str) -> bool:
    return len((text or "").strip()) < _MIN_TEXT_THRESHOLD


def _extract_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    chunks: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            chunks.append(page.extract_text() or "")
        except Exception as exc:  # pragma: no cover
            log.warning("pdf_page_extract_failed", page=i, error=str(exc))
    return "\n".join(chunks).strip()


def _extract_docx(path: Path) -> str:
    from docx import Document  # python-docx

    doc = Document(str(path))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts).strip()
